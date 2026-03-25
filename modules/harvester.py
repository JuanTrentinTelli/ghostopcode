"""
GhostOpcode harvester — crawl, file hunt, email/config harvest, document metadata.
"""

from __future__ import annotations

import datetime as dt
import hashlib
import random
import re
import threading
import time
from collections import deque
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

import requests
import urllib3
from bs4 import BeautifulSoup
from rich import box
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.text import Text

from config import DEFAULT_THREADS, DEFAULT_TIMEOUT, OUTPUT_DIR, USER_AGENT
from utils.output import display_findings
from utils.target_parser import Target

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore[misc, assignment]

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[misc, assignment]

try:
    import openpyxl
except ImportError:
    openpyxl = None  # type: ignore[misc, assignment]

try:
    import whois as pywhois
except ImportError:
    pywhois = None  # type: ignore[misc, assignment]

C_PRI = "#00FF41"
C_DIM = "#6F7F86"
C_ERR = "#FF3B3B"
C_WARN = "#E8C547"
C_MUTED = "#4A5A62"
C_ACCENT = "#8B9CA8"

console = Console(highlight=False, force_terminal=True)

MAX_DOWNLOAD_BYTES = 10 * 1024 * 1024
MAX_CRAWL_URLS_DEFAULT = 500

USER_AGENTS = [
    USER_AGENT,
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Safari/605.1.15",
]

FILE_EXTENSIONS: dict[str, dict[str, Any]] = {
    "pdf": {"risk": "HIGH", "extract_meta": True},
    "doc": {"risk": "HIGH", "extract_meta": True},
    "docx": {"risk": "HIGH", "extract_meta": True},
    "xls": {"risk": "HIGH", "extract_meta": True},
    "xlsx": {"risk": "HIGH", "extract_meta": True},
    "ppt": {"risk": "MEDIUM", "extract_meta": True},
    "pptx": {"risk": "MEDIUM", "extract_meta": True},
    "csv": {"risk": "HIGH", "extract_meta": False},
    "json": {"risk": "HIGH", "extract_meta": False},
    "xml": {"risk": "MEDIUM", "extract_meta": False},
    "sql": {"risk": "CRITICAL", "extract_meta": False},
    "env": {"risk": "CRITICAL", "extract_meta": False},
    "conf": {"risk": "HIGH", "extract_meta": False},
    "cfg": {"risk": "HIGH", "extract_meta": False},
    "ini": {"risk": "HIGH", "extract_meta": False},
    "yaml": {"risk": "HIGH", "extract_meta": False},
    "yml": {"risk": "HIGH", "extract_meta": False},
    "php": {"risk": "CRITICAL", "extract_meta": False},
    "py": {"risk": "CRITICAL", "extract_meta": False},
    "js": {"risk": "HIGH", "extract_meta": False},
    "sh": {"risk": "CRITICAL", "extract_meta": False},
    "bak": {"risk": "CRITICAL", "extract_meta": False},
    "zip": {"risk": "HIGH", "extract_meta": False},
    "tar": {"risk": "HIGH", "extract_meta": False},
    "gz": {"risk": "HIGH", "extract_meta": False},
}

_EXT_SET = set(FILE_EXTENSIONS.keys())

COMMON_FILE_PATHS: list[str] = [
    "backup.sql",
    "dump.sql",
    "database.sql",
    "db.sql",
    "backup.zip",
    "site.zip",
    "www.zip",
    "public.zip",
    ".env",
    ".env.local",
    ".env.production",
    ".env.backup",
    "wp-config.php",
    "config.php",
    "configuration.php",
    "database.yml",
    "database.yaml",
    "settings.py",
    "application.properties",
    "appsettings.json",
    "web.config",
    ".htaccess",
    "nginx.conf",
    "error.log",
    "access.log",
    "debug.log",
    "app.log",
    "laravel.log",
    "storage/logs/laravel.log",
    "phpinfo.php",
    "info.php",
    "test.php",
    "debug.php",
    "install.php",
    "setup.php",
    "upgrade.php",
    ".git/HEAD",
    ".git/config",
    ".gitignore",
    ".svn/entries",
    ".hg/hgrc",
    "sitemap.xml",
    "robots.txt",
    "crossdomain.xml",
    "clientaccesspolicy.xml",
    "security.txt",
    ".well-known/security.txt",
]

# Stricter email match: TLD 2–6 alpha, domain segment before TLD ≥ 2 chars (via validator)
EMAIL_REGEX = re.compile(
    r"[a-zA-Z0-9._%+\-]{2,}@[a-zA-Z0-9.\-]{3,}\.[a-zA-Z]{2,6}",
    re.IGNORECASE,
)

# Legacy obfuscated forms → normalize then validate with is_valid_email
EMAIL_OBFUSCATED = re.compile(
    r"[a-zA-Z0-9._%+\-]{2,}\s*[\[@]\s*[a-zA-Z0-9.\-]{3,}\s*[\.\u2024]\s*[a-zA-Z]{2,6}",
    re.IGNORECASE,
)


def is_valid_email(email: str) -> bool:
    """
    Validate extracted email beyond regex match.
    Filters obvious false positives (e.g. 7@w.vpe).
    """
    if not email or len(email) < 6:
        return False
    email = email.strip().lower()
    local, _, domain = email.rpartition("@")
    if not local or not domain:
        return False
    if len(local) < 2:
        return False
    if "." not in domain:
        return False
    tld = domain.rsplit(".", 1)[-1]
    if not tld.isalpha() or not (2 <= len(tld) <= 6):
        return False
    domain_body = domain.rsplit(".", 1)[0]
    if len(domain_body) < 2:
        return False
    if domain_body.replace(".", "").isdigit():
        return False
    # Single-letter local + single-letter host label (e.g. z@if.db) — regex may already exclude
    invalid_patterns = [
        r"^[a-z0-9]@[a-z0-9]\.[a-z]{2,6}$",
    ]
    for pattern in invalid_patterns:
        if re.search(pattern, email, re.IGNORECASE):
            return False
    return True


def _normalize_email_candidate(raw: str) -> str:
    """Collapse whitespace and common obfuscations for validation."""
    s = re.sub(r"\s+", "", raw.strip().replace("[at]", "@").replace("(at)", "@"))
    return s.lower()


def iter_valid_emails_in_text(text: str) -> list[str]:
    """All unique valid emails found in text using EMAIL_REGEX + obfuscated pattern."""
    seen: set[str] = set()
    out: list[str] = []
    if not text:
        return out
    for m in EMAIL_REGEX.finditer(text):
        em = _normalize_email_candidate(m.group(0))
        if is_valid_email(em) and em not in seen:
            seen.add(em)
            out.append(em)
    for m in EMAIL_OBFUSCATED.finditer(text):
        em = _normalize_email_candidate(m.group(0))
        if is_valid_email(em) and em not in seen:
            seen.add(em)
            out.append(em)
    return out

USERNAME_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("url_author", re.compile(r"author[=/]([a-zA-Z0-9_\-\.]+)", re.I)),
    ("url_profile", re.compile(r"profile[=/]([a-zA-Z0-9_\-\.]+)", re.I)),
    ("url_user", re.compile(r"user[=/]([a-zA-Z0-9_\-\.]+)", re.I)),
    ("byline", re.compile(r"by\s+([A-Z][a-z]+\s+[A-Z][a-z]+)")),
    ("json_author", re.compile(r'"author":\s*"([^"]+)"')),
]

LINKEDIN_PATTERNS: list[re.Pattern[str]] = [
    re.compile(r"linkedin\.com/in/([a-zA-Z0-9\-]+)", re.I),
    re.compile(r"linkedin\.com/pub/([a-zA-Z0-9\-/]+)", re.I),
]

ROLE_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^(ceo|diretor|director)", re.I), "C-Level / Director"),
    (re.compile(r"^(cto|ciso|cfo)", re.I), "C-Level"),
    (re.compile(r"^(ti\.|it\.|suporte|infra)", re.I), "IT / Infrastructure"),
    (re.compile(r"^(dev|developer|eng)", re.I), "Developer / Engineer"),
    (re.compile(r"^(rh\.|hr\.|people)", re.I), "HR / People"),
    (re.compile(r"^(comercial|vendas|sales)", re.I), "Sales / Commercial"),
    (re.compile(r"^(financeiro|finance)", re.I), "Finance"),
    (re.compile(r"^(marketing|mkt)", re.I), "Marketing"),
    (re.compile(r"^(contato|contact|info|admin|abuse)", re.I), "Generic / Admin"),
]

LEAK_SIGNATURES: dict[str, list[str]] = {
    "env": ["APP_KEY=", "DB_PASSWORD=", "SECRET_KEY=", "API_KEY=", "DATABASE_URL="],
    "git": ["ref: refs/heads/"],
    "wp-config": ["DB_NAME", "DB_PASSWORD", "AUTH_KEY"],
    "phpinfo": ["PHP Version", "php.ini", "SERVER_SOFTWARE"],
    "sql": ["INSERT INTO", "CREATE TABLE", "DROP TABLE"],
    "log": ["PHP Fatal", "Traceback", "Exception in"],
    "config.php": ["define(", "$db_", "password"],
    "laravel.log": ["local.ERROR", "production.ERROR"],
}

SECRET_PATTERNS: dict[str, list[re.Pattern[str]]] = {
    "database_password": [
        re.compile(r"DB_PASSWORD[=:]\s*['\"]?([^'\"\n\s#]+)", re.I),
        re.compile(r"database_password[=:]\s*['\"]?([^'\"\n\s#]+)", re.I),
        re.compile(r"db_pass[=:]\s*['\"]?([^'\"\n\s#]+)", re.I),
    ],
    "api_key": [
        re.compile(r"API_KEY[=:]\s*['\"]?([a-zA-Z0-9_\-]{16,})", re.I),
        re.compile(r"SECRET_KEY[=:]\s*['\"]?([a-zA-Z0-9_\-]{16,})", re.I),
        re.compile(r"APP_KEY[=:]\s*['\"]?([a-zA-Z0-9+/=]{20,})", re.I),
    ],
    "aws_key": [
        re.compile(r"(AKIA[0-9A-Z]{16})"),
        re.compile(r"aws_secret[=:]\s*['\"]?([^\s'\"]{20,})", re.I),
    ],
    "jwt_secret": [
        re.compile(r"JWT_SECRET[=:]\s*['\"]?([^\s'\"]{8,})", re.I),
    ],
    "private_key": [
        re.compile(r"-----BEGIN (RSA |EC )?PRIVATE KEY-----"),
    ],
}

_WIN_PATH_RE = re.compile(
    r"[A-Za-z]:\\(?:Users|Documents|Desktop|AppData)[^\"'\s<>]{3,200}",
    re.I,
)
_UNC_PATH_RE = re.compile(r"\\\\[A-Za-z0-9.\-]+\\[^\s\"'<>]{3,200}")


def mask_secret_display(value: str, visible: int = 8) -> str:
    """Mask secret for terminal (first N chars + asterisks)."""
    v = value.strip()
    if len(v) <= visible:
        return "*****"
    return v[:visible] + "*****"


def resolve_base_url(target: Target, timeout: float) -> str | None:
    """Pick working https or http origin for target host."""
    host = target.value
    for scheme in ("https", "http"):
        base = f"{scheme}://{host}".rstrip("/")
        root = base + "/"
        try:
            ua = random.choice(USER_AGENTS)
            r = requests.get(
                root,
                timeout=timeout,
                verify=False,
                allow_redirects=True,
                headers={"User-Agent": ua},
            )
            if r.status_code > 0 and r.status_code < 600:
                return urlparse(r.url)._replace(path="", params="", query="", fragment="").geturl().rstrip("/")
        except Exception:  # noqa: BLE001
            continue
    return None


def _normalize_netloc(netloc: str) -> str:
    n = netloc.lower()
    if n.startswith("www."):
        return n[4:]
    return n


def _same_site(url: str, base_netloc: str) -> bool:
    try:
        p = urlparse(url)
        if p.scheme not in ("http", "https") or not p.netloc:
            return False
        return _normalize_netloc(p.netloc) == _normalize_netloc(base_netloc)
    except Exception:  # noqa: BLE001
        return False


def _file_ext_from_url(url: str) -> str | None:
    path = urlparse(url).path.lower()
    if "." not in path:
        return None
    ext = path.rsplit(".", 1)[-1]
    if ext in _EXT_SET:
        return ext
    return None


def _session(timeout: float) -> requests.Session:
    s = requests.Session()
    s.verify = False
    s.headers.update({"User-Agent": USER_AGENT})
    s.request_timeout = timeout  # type: ignore[attr-defined]
    return s


def crawl(
    base_url: str,
    depth: int,
    timeout: float,
    max_urls: int,
    errors: list[str],
) -> tuple[list[dict[str, str]], set[str], int]:
    """
    BFS crawl same-site links; collect HTML pages and file URLs.
    Random delay 0.1–0.3s between requests; rotates User-Agent.
    Never raises.
    """
    pages: list[dict[str, str]] = []
    file_urls: set[str] = set()
    visited: set[str] = set()
    base_p = urlparse(base_url)
    base_netloc = base_p.netloc
    start = base_url.rstrip("/") + "/"
    q: deque[tuple[str, int]] = deque([(start, 0)])
    crawled = 0

    while q and len(visited) < max_urls:
        url, d = q.popleft()
        if url in visited:
            continue
        visited.add(url)
        if d > depth:
            continue

        time.sleep(random.uniform(0.1, 0.3))
        ua = random.choice(USER_AGENTS)
        try:
            r = requests.get(
                url,
                timeout=timeout,
                verify=False,
                allow_redirects=True,
                headers={"User-Agent": ua},
            )
            crawled += 1
            ct = (r.headers.get("content-type") or "").lower()
            final_url = r.url
            if not _same_site(final_url, base_netloc):
                continue

            ext_hit = _file_ext_from_url(final_url)
            if ext_hit and "text/html" not in ct:
                file_urls.add(final_url.split("#")[0])
                continue

            if "text/html" not in ct and "html" not in ct:
                if ext_hit:
                    file_urls.add(final_url.split("#")[0])
                continue

            html = r.text or ""
            pages.append({"url": final_url, "html": html})

            soup = BeautifulSoup(html, "html.parser")
            for tag in soup.find_all("a", href=True):
                href = (tag.get("href") or "").strip()
                if not href or href.startswith(("#", "javascript:", "mailto:", "tel:")):
                    continue
                abs_u = urljoin(final_url, href)
                abs_u = abs_u.split("#")[0]
                if not _same_site(abs_u, base_netloc):
                    continue
                fe = _file_ext_from_url(abs_u)
                if fe:
                    file_urls.add(abs_u)
                elif d < depth and abs_u not in visited:
                    q.append((abs_u, d + 1))

            for tag in soup.find_all(["script", "link"]):
                raw = tag.get("src") or tag.get("href") or ""
                if raw and not raw.startswith("data:"):
                    abs_u = urljoin(final_url, raw).split("#")[0]
                    if _same_site(abs_u, base_netloc):
                        fe = _file_ext_from_url(abs_u)
                        if fe:
                            file_urls.add(abs_u)
        except requests.RequestException as e:
            errors.append(f"crawl {url}: {e}")
        except Exception as e:  # noqa: BLE001
            errors.append(f"crawl {url}: {e}")

    return pages, file_urls, crawled


def brute_common_paths(
    base_url: str,
    paths: list[str],
    timeout: float,
    threads: int,
    errors: list[str],
) -> list[dict[str, Any]]:
    """GET known sensitive paths in parallel; return raw probe rows."""
    base = base_url.rstrip("/")
    results: list[dict[str, Any]] = []
    lock = threading.Lock()

    def probe(rel: str) -> None:
        url = f"{base}/{rel.lstrip('/')}"
        ua = random.choice(USER_AGENTS)
        try:
            time.sleep(random.uniform(0.05, 0.15))
            r = requests.get(
                url,
                timeout=timeout,
                verify=False,
                allow_redirects=True,
                headers={"User-Agent": ua},
            )
            body = (r.text or "")[:50000]
            with lock:
                results.append(
                    {
                        "path": "/" + rel.lstrip("/"),
                        "url": r.url,
                        "status": r.status_code,
                        "body": body,
                    }
                )
        except Exception as e:  # noqa: BLE001
            with lock:
                errors.append(f"brute {url}: {e}")

    with ThreadPoolExecutor(max_workers=max(1, threads)) as ex:
        futs = [ex.submit(probe, p) for p in paths]
        for f in as_completed(futs):
            try:
                f.result()
            except Exception as e:  # noqa: BLE001
                errors.append(str(e))

    return results


def validate_leak(path: str, status: int, body: str) -> dict[str, Any] | None:
    """Return confirmation dict if response looks like a real leak."""
    if status != 200 or not body.strip():
        return None
    pl = path.lower().replace("\\", "/")
    body_l = body[:50000]

    if ".env" in pl or pl.rstrip("/").endswith(".env"):
        if any(s in body_l for s in LEAK_SIGNATURES["env"]):
            return {"confirmed": True, "signature_key": "env"}
    if "wp-config" in pl:
        if any(s in body_l for s in LEAK_SIGNATURES["wp-config"]):
            return {"confirmed": True, "signature_key": "wp-config"}
    if ".git/head" in pl:
        if "ref:" in body_l.lower():
            return {"confirmed": True, "signature_key": "git"}
    if "phpinfo" in pl or pl.endswith("info.php"):
        if any(s in body_l for s in LEAK_SIGNATURES["phpinfo"]):
            return {"confirmed": True, "signature_key": "phpinfo"}
    if ".sql" in pl:
        if any(s in body_l for s in LEAK_SIGNATURES["sql"]):
            return {"confirmed": True, "signature_key": "sql"}
    if ".log" in pl or "/logs/" in pl:
        if any(s in body_l for s in LEAK_SIGNATURES["log"]):
            return {"confirmed": True, "signature_key": "log"}
    if "laravel.log" in pl:
        if any(s in body_l for s in LEAK_SIGNATURES["laravel.log"]):
            return {"confirmed": True, "signature_key": "laravel.log"}
    if pl.endswith("config.php") or "/config.php" in pl:
        if any(s in body_l for s in LEAK_SIGNATURES["config.php"]):
            return {"confirmed": True, "signature_key": "config.php"}
    return None


def extract_secrets(content: str, file_hint: str) -> list[dict[str, str]]:
    """Pull secret-like strings from leaked text (full values for JSON/report)."""
    found: list[dict[str, str]] = []
    for stype, patterns in SECRET_PATTERNS.items():
        for rx in patterns:
            for m in rx.finditer(content[:100000]):
                val = m.group(1) if m.lastindex else m.group(0)
                if val and len(val.strip()) > 2:
                    found.append({"type": stype, "value": val.strip()})
    # dedupe by type+value
    seen: set[tuple[str, str]] = set()
    out: list[dict[str, str]] = []
    for item in found:
        k = (item["type"], item["value"])
        if k not in seen:
            seen.add(k)
            out.append(item)
    return out[:40]


def harvest_emails_from_html(html: str, source_url: str) -> list[dict[str, Any]]:
    """Extract emails and usernames from HTML."""
    out: list[dict[str, Any]] = []
    if not html:
        return out

    for em in iter_valid_emails_in_text(html):
        out.append({"email": em, "source": source_url, "kind": "html_regex"})

    soup = BeautifulSoup(html, "html.parser")
    for a in soup.find_all("a", href=True):
        href = a.get("href", "")
        if href.lower().startswith("mailto:"):
            em = href[7:].split("?")[0].strip()
            em = _normalize_email_candidate(em.split(",")[0])
            if is_valid_email(em):
                out.append({"email": em, "source": source_url, "kind": "mailto"})

    for name, rx in USERNAME_PATTERNS:
        for m in rx.finditer(html):
            out.append(
                {
                    "username": m.group(1).strip(),
                    "source": source_url,
                    "kind": name,
                }
            )

    for rx in LINKEDIN_PATTERNS:
        for m in rx.finditer(html):
            out.append(
                {
                    "linkedin": m.group(1),
                    "source": source_url,
                    "kind": "linkedin",
                }
            )

    for script in soup.find_all("script", type="application/ld+json"):
        if script.string:
            for em in iter_valid_emails_in_text(script.string):
                out.append({"email": em, "source": source_url, "kind": "json-ld"})

    for m in re.finditer(r"<!--([\s\S]*?)-->", html):
        chunk = m.group(1)
        for em in iter_valid_emails_in_text(chunk):
            out.append(
                {
                    "email": em,
                    "source": source_url + " (comment)",
                    "kind": "comment",
                }
            )

    return out


def harvest_emails_from_file(file_path: str, file_type: str) -> list[dict[str, Any]]:
    """Extract emails from downloaded document text/metadata."""
    out: list[dict[str, Any]] = []
    try:
        if file_type == "pdf" and fitz is not None:
            doc = fitz.open(file_path)
            try:
                text = ""
                for i in range(min(doc.page_count, 30)):
                    text += doc.load_page(i).get_text() or ""
                for em in iter_valid_emails_in_text(text):
                    out.append(
                        {
                            "email": em,
                            "source": f"pdf:{Path(file_path).name}",
                            "kind": "pdf_text",
                        }
                    )
            finally:
                doc.close()
        elif file_type == "docx" and DocxDocument is not None:
            d = DocxDocument(file_path)
            for p in d.paragraphs[:500]:
                for em in iter_valid_emails_in_text(p.text):
                    out.append(
                        {
                            "email": em,
                            "source": f"docx:{Path(file_path).name}",
                            "kind": "docx_text",
                        }
                    )
        elif file_type == "xlsx" and openpyxl is not None:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            try:
                for sheet in wb.worksheets[:10]:
                    for row in sheet.iter_rows(max_row=200, max_col=30):
                        for cell in row:
                            v = cell.value
                            if isinstance(v, str):
                                for em in iter_valid_emails_in_text(v):
                                    out.append(
                                        {
                                            "email": em,
                                            "source": f"xlsx:{Path(file_path).name}",
                                            "kind": "xlsx_cell",
                                        }
                                    )
            finally:
                wb.close()
        else:
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")[:200000]
            for em in iter_valid_emails_in_text(raw):
                out.append(
                    {
                        "email": em,
                        "source": f"{file_type}:{Path(file_path).name}",
                        "kind": "text_file",
                    }
                )
    except Exception:  # noqa: BLE001
        pass
    return out


def enrich_email(email: str, domain: str) -> dict[str, Any]:
    """Infer role hint from local-part."""
    local = email.split("@", 1)[0].lower()
    role = "Unknown"
    for rx, label in ROLE_PATTERNS:
        if rx.search(local):
            role = label
            break
    name_guess = None
    if "." in local and "@" not in local:
        parts = re.split(r"[._]", local)
        if len(parts) >= 2 and all(p.isalpha() or len(p) > 1 for p in parts[:2]):
            name_guess = " ".join(p.capitalize() for p in parts[:2] if p.isalpha())
    return {"email": email, "role": role, "name": name_guess, "domain": domain}


def _clean_meta_str(val: object) -> str | None:
    """PyMuPDF often returns empty strings instead of None."""
    if val is None:
        return None
    s = str(val).strip()
    return s if s else None


def _normalize_pdf_date(val: str | None) -> str | None:
    """Turn PDF date strings (D:YYYYMMDDHHmmSS) into YYYY-MM-DD when possible."""
    s = _clean_meta_str(val)
    if not s:
        return None
    if s.startswith("D:"):
        digits = s[2:10]
        if len(digits) == 8 and digits.isdigit():
            return f"{digits[0:4]}-{digits[4:6]}-{digits[6:8]}"
    if len(s) >= 10 and s[4] == "-" and s[7] == "-":
        return s[:10]
    return s[:32]


def extract_pdf_metadata(file_path: str) -> dict[str, Any]:
    """
    Extract metadata from PDF using PyMuPDF.
    Returns empty/minimal dict on failure — never raises.
    """
    try:
        import fitz as _fitz  # PyMuPDF
    except ImportError:
        return {"error": "PyMuPDF not installed — pip install PyMuPDF"}

    try:
        doc = _fitz.open(file_path)
        try:
            raw_meta = doc.metadata or {}
            emails_in_doc: list[str] = []
            text_for_paths = ""
            n_pages = len(doc)
            for page_num in range(min(10, n_pages)):
                page = doc.load_page(page_num)
                text = page.get_text() or ""
                if page_num < 3:
                    for em in EMAIL_REGEX.findall(text):
                        el = em.lower()
                        if is_valid_email(el) and el not in emails_in_doc:
                            emails_in_doc.append(el)
                text_for_paths += text + "\n"
            internal_paths = list(
                dict.fromkeys(
                    _WIN_PATH_RE.findall(text_for_paths)
                    + _UNC_PATH_RE.findall(text_for_paths)
                )
            )[:20]
        finally:
            doc.close()

        author_s = _clean_meta_str(raw_meta.get("author", ""))
        creator_s = _clean_meta_str(raw_meta.get("creator", ""))
        producer_s = _clean_meta_str(raw_meta.get("producer", ""))

        result: dict[str, Any] = {
            "title": _clean_meta_str(raw_meta.get("title", "")),
            "author": author_s,
            "subject": _clean_meta_str(raw_meta.get("subject", "")),
            "keywords": _clean_meta_str(raw_meta.get("keywords", "")),
            "creator": creator_s,
            "producer": producer_s,
            "creation_date": _normalize_pdf_date(_clean_meta_str(raw_meta.get("creationDate", ""))),
            "mod_date": _normalize_pdf_date(_clean_meta_str(raw_meta.get("modDate", ""))),
            "emails_found": list(dict.fromkeys(emails_in_doc)),
            "internal_paths": internal_paths,
        }

        software = creator_s or producer_s
        if software:
            result["software"] = software

        if author_s:
            if "\\" in author_s:
                result["windows_username"] = author_s.split("\\")[-1].strip()
            elif "@" in author_s and is_valid_email(author_s.lower()):
                result["author_email"] = author_s.lower()

        return {k: v for k, v in result.items() if v not in (None, [], "")}
    except Exception as e:  # noqa: BLE001
        return {"error": str(e)}


def extract_docx_metadata(file_path: str) -> dict[str, Any]:
    """DOCX core properties."""
    out: dict[str, Any] = {
        "author": None,
        "last_modified_by": None,
        "created": None,
        "modified": None,
        "title": None,
        "subject": None,
        "keywords": None,
        "internal_paths": [],
        "emails_found": [],
        "software": None,
    }
    if DocxDocument is None:
        return out
    try:
        d = DocxDocument(file_path)
        cp = d.core_properties
        out["author"] = cp.author
        out["last_modified_by"] = cp.last_modified_by
        out["created"] = str(cp.created) if cp.created else None
        out["modified"] = str(cp.modified) if cp.modified else None
        out["title"] = cp.title
        out["subject"] = cp.subject
        out["keywords"] = cp.keywords
        text = "\n".join(p.text for p in d.paragraphs[:400])
        for em in iter_valid_emails_in_text(text):
            out["emails_found"].append(em)
        out["internal_paths"] = list(
            dict.fromkeys(_WIN_PATH_RE.findall(text) + _UNC_PATH_RE.findall(text))
        )[:20]
    except Exception:  # noqa: BLE001
        pass
    return out


def extract_xlsx_metadata(file_path: str) -> dict[str, Any]:
    """XLSX properties and sheet names."""
    out: dict[str, Any] = {
        "author": None,
        "created": None,
        "modified": None,
        "title": None,
        "sheet_names": [],
        "internal_paths": [],
        "emails_found": [],
        "software": None,
    }
    if openpyxl is None:
        return out
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        try:
            props = wb.properties
            out["author"] = props.creator
            out["created"] = str(props.created) if props.created else None
            out["modified"] = str(props.modified) if props.modified else None
            out["title"] = props.title
            out["sheet_names"] = list(wb.sheetnames)[:30]
            for sheet in wb.worksheets[:5]:
                for row in sheet.iter_rows(max_row=100, max_col=20):
                    for cell in row:
                        v = cell.value
                        if isinstance(v, str):
                            for em in iter_valid_emails_in_text(v):
                                out["emails_found"].append(em)
                            for p in _WIN_PATH_RE.findall(v) + _UNC_PATH_RE.findall(v):
                                if p not in out["internal_paths"]:
                                    out["internal_paths"].append(p)
        finally:
            wb.close()
        out["internal_paths"] = out["internal_paths"][:20]
    except Exception:  # noqa: BLE001
        pass
    return out


def correlate_metadata(
    all_metadata: list[dict[str, Any]],
    emails: list[dict[str, Any]],
    linkedin_profiles: list[str],
) -> dict[str, Any]:
    """Cross-reference document metadata, harvested emails, and LinkedIn handles."""
    usernames: set[str] = set()
    software: set[str] = set()
    internal_paths: set[str] = set()
    doc_emails: set[str] = set()
    dates: list[str] = []

    for m in all_metadata:
        if m.get("error"):
            continue
        for k in ("author", "last_modified_by", "windows_username"):
            v = m.get(k)
            if isinstance(v, str) and v.strip():
                usernames.add(v.strip())
        sw = m.get("software") or m.get("producer") or m.get("creator")
        if isinstance(sw, str) and sw.strip():
            software.add(sw.strip())
        for p in m.get("internal_paths") or []:
            internal_paths.add(p)
        for e in m.get("emails_found") or []:
            if isinstance(e, str) and e:
                doc_emails.add(e.lower())
        for e in m.get("emails_in_doc") or []:
            if isinstance(e, str) and e:
                doc_emails.add(e.lower())
        ae = m.get("author_email")
        if isinstance(ae, str) and ae:
            doc_emails.add(ae.lower())
        for dkey in ("creation_date", "mod_date", "created", "modified"):
            dv = m.get(dkey)
            if isinstance(dv, str) and len(dv) > 6:
                dates.append(dv[:10])

    for email_data in emails:
        addr = (email_data.get("email") or "").strip().lower()
        if not addr or addr == "—" or "@" not in addr:
            continue
        if not is_valid_email(addr):
            continue
        local = addr.split("@", 1)[0]
        if "." in local or "_" in local:
            usernames.add(local)

    for h in linkedin_profiles:
        hs = str(h).strip()
        if hs:
            usernames.add(hs)

    cr: tuple[str | None, str | None] = (None, None)
    if dates:
        cr = (min(dates), max(dates))

    ad_guess = None
    for p in internal_paths:
        if "\\\\" in p and "\\" in p:
            parts = p.split("\\")
            if len(parts) > 2:
                ad_guess = parts[2].upper() if parts[2] else None
                break

    return {
        "usernames": sorted(usernames),
        "software_stack": sorted(software),
        "internal_paths": sorted(internal_paths)[:30],
        "emails_in_docs": sorted(doc_emails),
        "creation_range": cr,
        "ad_domain_guess": ad_guess,
    }


def _pdf_meta_meaningful(md: dict[str, Any]) -> bool:
    """True if PDF metadata / extracted fields contain any useful signal."""
    keys = (
        "author",
        "creator",
        "producer",
        "creation_date",
        "mod_date",
        "created",
        "modified",
        "title",
        "subject",
        "keywords",
        "software",
        "windows_username",
        "author_email",
        "last_modified_by",
    )
    for k in keys:
        v = md.get(k)
        if isinstance(v, str) and v.strip():
            return True
    for k in ("emails_found", "emails_in_doc", "internal_paths", "sheet_names"):
        v = md.get(k)
        if isinstance(v, list) and len(v) > 0:
            return True
    return False


def _human_size(n: int) -> str:
    if n >= 1024 * 1024:
        return f"{n / (1024 * 1024):.1f}MB"
    if n >= 1024:
        return f"{n / 1024:.1f}KB"
    return f"{n}b"


def _whois_emails(domain: str, timeout: int, errors: list[str]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    if pywhois is None:
        return out
    try:
        import socket

        socket.setdefaulttimeout(float(timeout))
        w = pywhois.whois(domain)
        socket.setdefaulttimeout(None)
        em = getattr(w, "emails", None)
        if isinstance(em, str):
            lst = [em]
        elif isinstance(em, list):
            lst = [str(x) for x in em if x]
        else:
            lst = []
        for e in lst:
            el = e.lower().strip()
            if is_valid_email(el):
                out.append(
                    {
                        "email": el,
                        "source": "WHOIS",
                        "kind": "whois",
                    }
                )
    except Exception as e:  # noqa: BLE001
        errors.append(f"whois emails: {e}")
    return out


def run(target: Target, config: dict[str, Any]) -> dict[str, Any]:
    """
    Run four harvest phases: crawl, file/brute probes, email merge, metadata.
    Never raises.
    """
    t0 = time.perf_counter()
    threads = max(1, int(config.get("threads") or DEFAULT_THREADS))
    timeout = max(1.0, float(config.get("timeout") or DEFAULT_TIMEOUT))
    depth = max(0, int(config.get("depth", 3)))
    save_files = bool(config.get("save_files", True))
    verbose = bool(config.get("verbose", False))
    max_crawl = int(config.get("max_crawl_urls") or MAX_CRAWL_URLS_DEFAULT)

    errors: list[str] = []
    base: dict[str, Any] = {
        "module": "harvester",
        "target": target.value,
        "status": "pending",
        "files": [],
        "emails": [],
        "config_leaks": [],
        "intelligence": {},
        "stats": {
            "urls_crawled": 0,
            "files_found": 0,
            "emails_found": 0,
            "leaks_found": 0,
            "duration_s": 0.0,
        },
        "errors": errors,
        "findings": [],
    }

    if target.is_cidr():
        console.print(
            Panel(
                Text("  HARVESTER  ·  CIDR not supported", style=f"bold {C_PRI}"),
                border_style=C_ACCENT,
                box=box.DOUBLE,
            )
        )
        console.print(Text("  [SKIP] Harvester — domain or IP only.", style=C_WARN))
        base["status"] = "skipped"
        return base

    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = re.sub(r"[^\w.\-]+", "_", target.value, flags=re.ASCII)[:80] or "target"
    default_out = Path(OUTPUT_DIR) / f"{slug}_{stamp}" / "files"
    out_dir = Path(config.get("output_dir") or str(default_out))
    if save_files:
        try:
            out_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            errors.append(str(e))
            save_files = False

    try:
        base_url = resolve_base_url(target, timeout)
        if not base_url:
            base["status"] = "error"
            base["error"] = "Could not reach target (HTTP/HTTPS)"
            errors.append(base["error"])
            console.print(Text(f"  [✗] {base['error']}", style=C_ERR))
            return base

        console.print(
            Panel(
                Text(
                    f"  HARVESTER  ·  {target.value}  ·  files + emails + leaks",
                    style=f"bold {C_PRI}",
                ),
                border_style=C_ACCENT,
                box=box.DOUBLE,
                width=min(console.size.width, 82) if console.size else 82,
            )
        )

        interrupted = False
        pages: list[dict[str, str]] = []
        file_urls: set[str] = set()
        crawled_n = 0
        brute_rows: list[dict[str, Any]] = []
        config_leaks: list[dict[str, Any]] = []
        files_out: list[dict[str, Any]] = []
        all_meta_raw: list[dict[str, Any]] = []
        email_map: dict[str, dict[str, Any]] = {}

        # --- Phase 1 crawl ---
        console.print(
            Text(
                f"\n [1/4] Crawling {base_url} (depth: {depth})...",
                style=f"bold {C_WARN}",
            )
        )
        try:
            pages, file_urls, crawled_n = crawl(
                base_url, depth, timeout, max_crawl, errors
            )
        except KeyboardInterrupt:
            interrupted = True
            errors.append("[!] Interrupted during crawl — partial data")

        t1 = time.perf_counter()
        console.print(
            Text(
                f"       {len(pages)} pages · {len(file_urls)} file URLs · "
                f"{crawled_n} fetches in {t1 - t0:.1f}s",
                style=C_DIM,
            )
        )

        # --- Phase 2 brute paths + classify leaks ---
        console.print(
            Text("\n [2/4] File hunt + config path brute...", style=f"bold {C_WARN}")
        )
        try:
            brute_rows = brute_common_paths(
                base_url, COMMON_FILE_PATHS, timeout, threads, errors
            )
        except KeyboardInterrupt:
            interrupted = True
            errors.append("[!] Interrupted during brute")

        for row in brute_rows:
            path = row["path"]
            body = row.get("body") or ""
            val = validate_leak(path, int(row.get("status", 0)), body)
            if val and val.get("confirmed"):
                secrets = extract_secrets(body, path)
                config_leaks.append(
                    {
                        "path": path,
                        "url": row.get("url"),
                        "status": row.get("status"),
                        "confirmed": True,
                        "secrets": secrets,
                        "signature_key": val.get("signature_key") or "config_leak",
                    }
                )
                # surface file URL for download consideration
                if _file_ext_from_url(row["url"]) or any(
                    path.endswith(x) for x in (".sql", ".zip", ".env", ".log")
                ):
                    file_urls.add(row["url"].split("#")[0])

        if config_leaks:
            display_findings(
                [
                    {
                        "risk": "CRITICAL",
                        "category": str(leak.get("signature_key") or "config_leak"),
                        "value": str(leak.get("url") or leak.get("path") or ""),
                        "note": (
                            f"Confirmed leak — {len(leak.get('secrets') or [])} secret(s) · "
                            f"{leak.get('path', '')}"
                        ),
                    }
                    for leak in config_leaks
                    if leak.get("confirmed")
                ],
                module="harvester",
                verbose=bool(config.get("verbose")),
            )
        for leak in config_leaks:
            for s in leak.get("secrets", [])[:6]:
                console.print(
                    Text(
                        f"       ├── {s['type']}: {mask_secret_display(s['value'])}",
                        style=C_ERR,
                    )
                )

        # Merge brute-found interesting URLs into file set
        for row in brute_rows:
            if int(row.get("status", 0)) == 200 and _file_ext_from_url(row["url"]):
                file_urls.add(row["url"].split("#")[0])

        console.print(
            Text(
                f"       {len(file_urls)} unique file URLs · "
                f"{len(config_leaks)} leak(s) confirmed",
                style=C_DIM,
            )
        )

        # --- Download files + harvest text from disk ---
        console.print(Text("\n [3/4] Downloading + harvesting text...", style=f"bold {C_WARN}"))
        try:
            for furl in sorted(file_urls):
                ext = _file_ext_from_url(furl)
                if not ext:
                    continue
                spec = FILE_EXTENSIONS.get(ext, {"risk": "MEDIUM", "extract_meta": False})
                fn = Path(urlparse(furl).path).name or "download"
                local_path: str | None = None
                size_b = 0
                downloaded = False
                try:
                    ua = random.choice(USER_AGENTS)
                    head = requests.head(
                        furl,
                        timeout=timeout,
                        verify=False,
                        allow_redirects=True,
                        headers={"User-Agent": ua},
                    )
                    cl = head.headers.get("Content-Length")
                    if cl and cl.isdigit() and int(cl) > MAX_DOWNLOAD_BYTES:
                        files_out.append(
                            {
                                "url": furl,
                                "filename": fn,
                                "type": ext,
                                "size": _human_size(int(cl)),
                                "size_bytes": int(cl),
                                "risk": spec["risk"],
                                "downloaded": False,
                                "local_path": None,
                                "metadata": {},
                                "note": "skipped (>10MB)",
                            }
                        )
                        continue
                    r = requests.get(
                        furl,
                        timeout=timeout,
                        verify=False,
                        allow_redirects=True,
                        headers={"User-Agent": ua},
                        stream=True,
                    )
                    data = b""
                    for chunk in r.iter_content(65536):
                        data += chunk
                        if len(data) > MAX_DOWNLOAD_BYTES:
                            break
                    size_b = len(data)
                    if save_files and size_b <= MAX_DOWNLOAD_BYTES:
                        safe = hashlib.sha256(furl.encode()).hexdigest()[:12] + "_" + fn
                        lp = out_dir / safe
                        lp.write_bytes(data)
                        local_path = str(lp.resolve())
                        downloaded = True

                    meta: dict[str, Any] = {}
                    if (
                        downloaded
                        and local_path
                        and spec.get("extract_meta")
                        and size_b < MAX_DOWNLOAD_BYTES
                    ):
                        if ext == "pdf":
                            meta = extract_pdf_metadata(local_path)
                        elif ext == "docx":
                            meta = extract_docx_metadata(local_path)
                        elif ext == "xlsx":
                            meta = extract_xlsx_metadata(local_path)
                        meta = {k: v for k, v in meta.items() if v not in (None, [], "")}
                        all_meta_raw.append(meta)

                    files_out.append(
                        {
                            "url": furl,
                            "filename": fn,
                            "type": ext,
                            "size": _human_size(size_b),
                            "size_bytes": size_b,
                            "risk": spec["risk"],
                            "downloaded": downloaded,
                            "local_path": local_path,
                            "metadata": meta,
                        }
                    )
                    if downloaded:
                        console.print(
                            Text(
                                f"   [{spec['risk']:<8}] {fn[:40]:<42} {_human_size(size_b):>8}",
                                style=C_PRI if spec["risk"] != "CRITICAL" else C_ERR,
                            )
                        )
                except Exception as e:  # noqa: BLE001
                    errors.append(f"download {furl}: {e}")
                    if verbose:
                        console.print(Text(f"   [!] {furl} — {e}", style=C_WARN))
        except KeyboardInterrupt:
            interrupted = True
            errors.append("[!] Interrupted during file downloads — partial files kept")

        # --- Emails from HTML + files + WHOIS ---
        linkedin_profiles: list[str] = []

        def add_email_record(rec: dict[str, Any]) -> None:
            if "email" not in rec:
                return
            em = rec["email"].lower().strip()
            if "@" not in em:
                return
            if not is_valid_email(em):
                return
            dom = em.split("@", 1)[-1]
            enriched = enrich_email(em, dom)
            if em not in email_map:
                email_map[em] = {
                    "email": em,
                    "source": rec.get("source", ""),
                    "name": enriched.get("name"),
                    "role": enriched["role"],
                    "linkedin": None,
                }
            else:
                src = email_map[em]["source"]
                new_src = rec.get("source", "")
                if new_src and new_src not in src:
                    email_map[em]["source"] = f"{src}; {new_src}" if src else new_src

        try:
            for p in pages:
                for rec in harvest_emails_from_html(p["html"], p["url"]):
                    if "email" in rec:
                        add_email_record(rec)
                    if rec.get("linkedin"):
                        h = str(rec["linkedin"]).strip()
                        if h:
                            linkedin_profiles.append(h)
                        key = f"linkedin:{h}"
                        if h and key not in email_map:
                            email_map[key] = {
                                "email": "—",
                                "source": rec.get("source", ""),
                                "name": h,
                                "role": "LinkedIn",
                                "linkedin": f"https://linkedin.com/in/{h}",
                            }

            for fo in files_out:
                if fo.get("downloaded") and fo.get("local_path"):
                    for rec in harvest_emails_from_file(
                        fo["local_path"],
                        fo.get("type") or "",
                    ):
                        add_email_record(rec)

            if target.is_domain():
                for rec in _whois_emails(target.value, int(timeout), errors):
                    add_email_record(rec)
        except KeyboardInterrupt:
            interrupted = True
            errors.append("[!] Interrupted during email harvest — partial identities kept")

        console.print(Text("\n [4/4] Metadata correlation + email table...", style=f"bold {C_WARN}"))

        emails_list = [v for k, v in email_map.items() if not k.startswith("linkedin:")]
        emails_list += [v for k, v in email_map.items() if k.startswith("linkedin:")]
        linkedin_unique = list(dict.fromkeys(linkedin_profiles))

        intel = correlate_metadata(all_meta_raw, emails_list, linkedin_unique)
        intel["emails_total"] = len([k for k in email_map if "@" in k])
        intel["files_total"] = len(files_out)
        intel["leaks_total"] = len(config_leaks)

        for fo in files_out:
            if not fo.get("downloaded"):
                continue
            md = fo.get("metadata") or {}
            ft = fo.get("type") or ""
            if ft == "pdf":
                console.print(Text(f"   {fo['filename']}", style=C_DIM))
                if md.get("error"):
                    console.print(
                        Text(f"       └── [i] {md['error']}", style=C_WARN),
                    )
                elif not _pdf_meta_meaningful(md):
                    console.print(
                        Text(
                            "       └── [i] No metadata found (PDF may be scanned/image-only)",
                            style=C_MUTED,
                        ),
                    )
                else:
                    for label, key in (
                        ("Author", "author"),
                        ("Creator", "creator"),
                        ("Producer", "producer"),
                        ("Created", "creation_date"),
                    ):
                        v = md.get(key)
                        if v and str(v).strip():
                            console.print(
                                Text(
                                    f"       ├── {label:<9}: {str(v)[:75]}",
                                    style=C_MUTED,
                                ),
                            )
                    ef = md.get("emails_found") or md.get("emails_in_doc")
                    if ef:
                        es = ", ".join(ef) if isinstance(ef, list) else str(ef)
                        console.print(
                            Text(f"       └── Emails   : {es[:200]}", style=C_MUTED),
                        )
            elif ft in ("docx", "xlsx") and _pdf_meta_meaningful(md):
                console.print(Text(f"   {fo['filename']}", style=C_DIM))
                for label, key in (
                    ("Author", "author"),
                    ("Last modified", "last_modified_by"),
                    ("Software", "software"),
                    ("Internal path", "internal_paths"),
                ):
                    v = md.get(key)
                    if v:
                        if isinstance(v, list):
                            v = v[0] if v else ""
                        console.print(
                            Text(f"       ├── {label}: {str(v)[:70]}", style=C_MUTED),
                        )

        # Email table
        if emails_list:
            tbl = Table(
                title=Text("Emails / identities", style=f"bold {C_PRI}"),
                box=box.ROUNDED,
                border_style=C_ACCENT,
            )
            tbl.add_column("Email / ID", style=C_DIM)
            tbl.add_column("Name", style=C_MUTED)
            tbl.add_column("Role (inferred)", style=C_DIM)
            for row in emails_list[:40]:
                tbl.add_row(
                    row.get("email") or row.get("linkedin") or "—",
                    row.get("name") or "—",
                    row.get("role") or "—",
                )
            console.print()
            console.print(tbl)

        console.print()
        console.print(Text(" [INTEL] Correlated signals", style=f"bold {C_PRI}"))
        console.print(
            Text(
                f"   ├── Usernames     : {' · '.join(intel.get('usernames') or ['—'])}",
                style=C_DIM,
            )
        )
        console.print(
            Text(
                f"   ├── Internal path : {' · '.join((intel.get('internal_paths') or ['—'])[:3])}",
                style=C_DIM,
            )
        )
        console.print(
            Text(
                f"   ├── Software      : {' · '.join((intel.get('software_stack') or ['—'])[:4])}",
                style=C_DIM,
            )
        )
        if intel.get("ad_domain_guess"):
            console.print(
                Text(
                    f"   └── AD guess      : {intel['ad_domain_guess']}",
                    style=C_WARN,
                )
            )

        duration = time.perf_counter() - t0
        unique_domains = {
            e.split("@")[-1] for e in email_map if "@" in e and not e.startswith("linkedin:")
        }

        base["status"] = "success"
        if interrupted:
            console.print(
                Text(
                    "  [!] Run interrupted (CTRL+C) — partial results above",
                    style=C_WARN,
                )
            )
        base["files"] = files_out
        base["emails"] = emails_list
        base["config_leaks"] = config_leaks
        base["intelligence"] = dict(intel)
        base["intelligence"]["linkedin_profiles"] = linkedin_unique
        base["stats"]["urls_crawled"] = crawled_n
        base["stats"]["files_found"] = len(files_out)
        base["stats"]["emails_found"] = len(email_map)
        base["stats"]["leaks_found"] = len(config_leaks)
        base["stats"]["duration_s"] = round(duration, 2)
        base["findings"] = files_out + config_leaks

        console.print()
        console.print(
            Text.assemble(
                ("\n [✓] Harvester complete\n", f"bold {C_PRI}"),
                (f"     URLs crawled  : {crawled_n}\n", C_DIM),
                (
                    f"     Files found   : {len(files_out)}  "
                    f"({sum(1 for f in files_out if f.get('downloaded'))} downloaded)\n",
                    C_DIM,
                ),
                (
                    f"     Emails found  : {len(email_map)}  "
                    f"({len(unique_domains)} domains)\n",
                    C_DIM,
                ),
                (
                    f"     Config leaks  : {len(config_leaks)}  "
                    f"({sum(1 for l in config_leaks if l.get('confirmed'))} confirmed)\n",
                    C_DIM,
                ),
                (
                    f"     Usernames     : {len(intel.get('usernames') or [])} from metadata\n",
                    C_DIM,
                ),
                (f"     Duration      : {duration:.1f}s", C_DIM),
            )
        )

    except Exception as e:  # noqa: BLE001
        base["status"] = "error"
        base["error"] = str(e)
        errors.append(str(e))
        console.print(Text(f"  [✗] {e}", style=C_ERR))

    return base
